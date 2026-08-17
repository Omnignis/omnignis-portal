#!/usr/bin/env python3
"""
Report exporters for the Omnignis livestream report.

One builder per format. Each takes (church_name, rows, period_label) and
returns raw bytes, so report.py can attach any combination to one email.

rows is a list of dicts: {date, title, total_views, unique_viewers}
"""

import io
import csv

import brand

# Column catalogue. Only thresholds Facebook actually publishes: the ladder on
# video_insights stops at 60 seconds, so there is no 5 or 10 minute option.
COLUMNS = {
    "date":       {"label": "Livestream Date", "width": 16, "numeric": False},
    "title":      {"label": "Title",           "width": 42, "numeric": False},
    "views":      {"label": "Total Views",     "width": 13, "numeric": True},
    "unique":     {"label": "Unique Viewers",  "width": 15, "numeric": True},
    "sec10":      {"label": "10s+ Viewers",    "width": 13, "numeric": True},
    "min1":       {"label": "1 min+ Viewers",  "width": 15, "numeric": True},
    "day_of":     {"label": "On Service Day",  "width": 15, "numeric": True},
    "week":       {"label": "During the Week", "width": 16, "numeric": True},
    "window":     {"label": "Through Saturday","width": 16, "numeric": True},
}

ROW_KEY = {
    "date": "date", "title": "title",
    "views": "total_views", "unique": "unique_viewers",
    "sec10": "sec10_viewers", "min1": "min1_viewers",
    "day_of": "day_of", "week": "during_week", "window": "through_saturday",
}

# What a church gets unless it turns on custom reporting.
DEFAULT_COLUMNS = ["date", "title", "views", "unique", "min1"]
SELECTABLE = ["views", "unique", "sec10", "min1", "day_of", "week", "window"]

# Shown when a figure needs snapshot history we do not have yet. Better than a
# zero, which a church would read as "nobody watched".
PENDING = "collecting"


def parse_columns(raw, custom=False):
    """Resolve a stored preference into an ordered column list."""
    if not custom:
        return list(DEFAULT_COLUMNS)
    wanted = {c.strip().lower() for c in (raw or "").split(",") if c.strip()}
    picked = [c for c in SELECTABLE if c in wanted]
    return ["date", "title"] + (picked or ["views", "unique", "min1"])


def cell_value(row, col):
    v = row.get(ROW_KEY[col])
    if v is None:
        return PENDING if col in ("day_of", "week", "sec10", "min1", "window") else ""
    return v


def headers_for(columns):
    return [COLUMNS[c]["label"] for c in columns]


def row_values(row, columns):
    return [cell_value(row, c) for c in columns]


def totals_row(rows, columns):
    """Total line. Only sums columns that are actually numeric and present."""
    out = []
    for i, c in enumerate(columns):
        if c == "title":
            out.append("Total")
        elif COLUMNS[c]["numeric"]:
            vals = [row.get(ROW_KEY[c]) for row in rows]
            nums = [v for v in vals if isinstance(v, (int, float))]
            out.append(sum(nums) if nums else "")
        else:
            out.append("")
    return out


HEADERS = ["Livestream Date", "Title", "Total Views", "Unique Viewers"]   # legacy default

EMBER = "FF6A1A"
EMBER_HEX = "#ff6a1a"
DARK_HEX = "#0d0f14"

# Formats offered to churches, in the order they appear in the UI.
SUPPORTED = ["xlsx", "pdf", "csv", "docx", "txt", "png"]

EXTENSION = {
    "xlsx": "xlsx", "pdf": "pdf", "csv": "csv",
    "docx": "docx", "txt": "txt", "png": "png",
}

LABEL = {
    "xlsx": "Excel", "pdf": "PDF", "csv": "CSV",
    "docx": "Word", "txt": "Plain text", "png": "Image (PNG)",
}


# --------------------------------------------------------------------------
# Spreadsheet formula injection guard.
#
# A video titled "=HYPERLINK(...)" or "+1-555-..." is a live formula the moment
# Excel or Sheets opens the file. Church staff open these attachments without
# thinking about it. Prefixing with an apostrophe makes the cell text.
# Only strings are touched, so real numbers are never mangled.
# --------------------------------------------------------------------------
_RISKY_PREFIXES = ("=", "+", "-", "@", "\t", "\r")


def sanitize_cell(value):
    if isinstance(value, str) and value[:1] in _RISKY_PREFIXES:
        return "'" + value
    return value


def _totals(rows):
    return (
        sum(r.get("total_views") or 0 for r in rows),
        sum(r.get("unique_viewers") or 0 for r in rows),
    )


# ----------------------------- xlsx -----------------------------
def build_xlsx(church_name, rows, period_label, columns=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    cols = columns or DEFAULT_COLUMNS
    n = len(cols)
    wb = Workbook()
    ws = wb.active
    ws.title = "Livestream Report"
    thin = Side(style="thin", color="DDDDDD")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    last_col = get_column_letter(n)
    ws.merge_cells(f"A1:{last_col}1")
    ws["A1"] = sanitize_cell(church_name)
    ws["A1"].font = Font(size=16, bold=True, color="0D0F14")
    ws.merge_cells(f"A2:{last_col}2")
    ws["A2"] = f"Livestream attendance report  |  {period_label}"
    ws["A2"].font = Font(size=10, italic=True, color="666666")

    ws.append([])
    ws.append(headers_for(cols))
    hr = ws.max_row
    for c in range(1, n + 1):
        cell = ws.cell(row=hr, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=EMBER)
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border

    for row in rows:
        ws.append([sanitize_cell(v) for v in row_values(row, cols)])
        for c in range(1, n + 1):
            ws.cell(row=ws.max_row, column=c).border = border

    ws.append(totals_row(rows, cols))
    for c in range(1, n + 1):
        cell = ws.cell(row=ws.max_row, column=c)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="F0F0F0")
        cell.border = border

    for i, key in enumerate(cols, start=1):
        ws.column_dimensions[get_column_letter(i)].width = COLUMNS[key]["width"] + 4

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ----------------------------- csv ------------------------------
def build_csv(church_name, rows, period_label, columns=None):
    cols = columns or DEFAULT_COLUMNS
    buf = io.StringIO()
    w = csv.writer(buf, quoting=csv.QUOTE_MINIMAL, lineterminator="\r\n")
    w.writerow([sanitize_cell(church_name)])
    w.writerow([f"Livestream attendance report | {period_label}"])
    w.writerow([])
    w.writerow(headers_for(cols))
    for r in rows:
        w.writerow([sanitize_cell(v) for v in row_values(r, cols)])
    w.writerow(totals_row(rows, cols))
    return buf.getvalue().encode("utf-8-sig")


# ----------------------------- txt ------------------------------
def build_txt(church_name, rows, period_label, columns=None):
    cols = columns or DEFAULT_COLUMNS
    widths = [COLUMNS[c]["width"] for c in cols]

    def line(cells):
        return " ".join(str(c)[:w].ljust(w) for c, w in zip(cells, widths)).rstrip()

    rule = "-" * (sum(widths) + len(widths) - 1)
    out = [church_name, f"Livestream attendance report | {period_label}", ""]
    out.append(line(headers_for(cols)))
    out.append(rule)
    for r in rows:
        out.append(line(row_values(r, cols)))
    out.append(rule)
    out.append(line(totals_row(rows, cols)))
    out.append("")
    out.append(brand.FOOTER_TEXT)
    return ("\n".join(out) + "\n").encode("utf-8")


# ----------------------------- pdf ------------------------------
def build_pdf(church_name, rows, period_label, columns=None):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

    cols_probe = columns or DEFAULT_COLUMNS
    # Past six columns a portrait page forces the headers to clip, so turn the
    # sheet sideways instead of shrinking the text into illegibility.
    wide = len(cols_probe) > 6
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(letter) if wide else letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        title=f"{church_name} livestream report",
    )
    styles = getSampleStyleSheet()
    h = ParagraphStyle("H", parent=styles["Heading1"], fontSize=17, leading=21,
                       textColor=colors.HexColor(DARK_HEX), spaceAfter=2)
    sub = ParagraphStyle("S", parent=styles["Normal"], fontSize=9.5,
                         textColor=colors.HexColor("#666666"), spaceAfter=16)
    cell = ParagraphStyle("C", parent=styles["Normal"], fontSize=9, leading=12)

    cols = columns or DEFAULT_COLUMNS

    def fmt(v):
        return f"{v:,}" if isinstance(v, (int, float)) else str(v)

    # Headers as Paragraphs so long labels wrap rather than being cut off.
    head_style = ParagraphStyle("TH", parent=styles["Normal"], fontSize=8 if wide else 8.5,
                                leading=10, alignment=1, textColor=colors.white,
                                fontName="Helvetica-Bold")
    data = [[Paragraph(h, head_style) for h in headers_for(cols)]]
    for r in rows:
        data.append([
            Paragraph(str(cell_value(r, c)), cell) if c == "title" else fmt(cell_value(r, c))
            for c in cols
        ])
    data.append([fmt(v) if v != "" else "" for v in totals_row(rows, cols)])

    # Give Title the slack and split the rest evenly, so 5 or 8 columns both fit.
    usable = 10.0 if wide else 7.1
    fixed = {"date": 1.05, "title": None}
    others = [c for c in cols if c not in ("date", "title")]
    per = max(0.72, (usable - 1.05 - 1.5) / max(len(others), 1))
    widths = []
    for c in cols:
        if c == "date":
            widths.append(1.05 * inch)
        elif c == "title":
            widths.append(max(1.5, usable - 1.05 - per * len(others)) * inch)
        else:
            widths.append(per * inch)

    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(EMBER_HEX)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8 if len(cols) > 6 else 9),
        ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDDDDD")),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F0F0F0")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#FAFAFA")]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))

    from reportlab.platypus import Image as RLImage

    # Footer: "Generated by Omnignis LLC" with the flame mark to its right.
    mark = RLImage(brand.flame_stream(), width=9, height=11.5)
    footer_style = ParagraphStyle("F", parent=styles["Normal"], fontSize=8.5,
                                  textColor=colors.HexColor("#666666"))
    footer = Table(
        [[Paragraph(brand.FOOTER_TEXT, footer_style), mark]],
        colWidths=[None, 16], hAlign="LEFT",
    )
    footer.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (0, 0), 5),
        ("RIGHTPADDING", (1, 0), (1, 0), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    doc.build([
        Paragraph(church_name, h),
        Paragraph(f"Livestream attendance report &nbsp;|&nbsp; {period_label}", sub),
        table, Spacer(1, 20),
        footer,
    ])
    return buf.getvalue()


# ----------------------------- docx -----------------------------
def build_docx(church_name, rows, period_label, columns=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade(cell, hex_colour):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hex_colour)
        cell._tc.get_or_add_tcPr().append(el)

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run(church_name)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0x0D, 0x0F, 0x14)

    meta = doc.add_paragraph()
    mrun = meta.add_run(f"Livestream attendance report  |  {period_label}")
    mrun.italic = True
    mrun.font.size = Pt(9.5)
    mrun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    cols = columns or DEFAULT_COLUMNS
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, head in enumerate(headers_for(cols)):
        c = table.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(head)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, EMBER)

    def fmt(v):
        return f"{v:,}" if isinstance(v, (int, float)) else str(v)

    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_values(row, cols)):
            cells[i].text = fmt(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    cells = table.add_row().cells
    for i, v in enumerate([fmt(v) if v != "" else "" for v in totals_row(rows, cols)]):
        cells[i].text = ""
        r = cells[i].paragraphs[0].add_run(str(v))
        r.bold = True
        r.font.size = Pt(9)
        shade(cells[i], "F0F0F0")

    footer = doc.add_paragraph()
    frun = footer.add_run(brand.FOOTER_TEXT + "  ")
    frun.italic = True
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    footer.add_run().add_picture(brand.flame_stream(), height=Pt(11))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------- png ------------------------------
def build_png(church_name, rows, period_label, columns=None):
    import matplotlib
    matplotlib.use("Agg")                       # no display in CI
    import matplotlib.pyplot as plt

    cols = columns or DEFAULT_COLUMNS

    def fmt(v):
        return f"{v:,}" if isinstance(v, (int, float)) else str(v)

    body = [[str(cell_value(r, c))[:52] if c == "title" else fmt(cell_value(r, c))
             for c in cols] for r in rows]
    body.append([fmt(v) if v != "" else "" for v in totals_row(rows, cols)])

    # Explicit geometry. A full-figure axes with loc="upper center" made the
    # title overlap the table and left dead space underneath, so the header
    # band and the table each get their own reserved height.
    header_in = 1.0
    row_in = 0.36
    n_rows = len(body) + 1                      # + the header row
    table_in = row_in * n_rows
    # Widen when there are more columns so nothing has to be squeezed.
    fig_w = 10.5 + max(0, len(cols) - 5) * 1.15
    fig_h = header_in + table_in + 0.22

    fig = plt.figure(figsize=(fig_w, fig_h), dpi=170)
    fig.patch.set_facecolor("white")

    fig.text(0.014, 1 - 0.30 / fig_h, church_name,
             fontsize=16, fontweight="bold", color=DARK_HEX, va="top")
    fig.text(0.014, 1 - 0.72 / fig_h,
             f"Livestream attendance report  |  {period_label}",
             fontsize=9.5, color="#666666", va="top", style="italic")

    ax = fig.add_axes([0.014, 0.16 / fig_h, 0.972, table_in / fig_h])
    ax.axis("off")

    # bbox makes the table fill the axes exactly, so row heights are even and
    # nothing needs scale() guesswork.
    weights = [COLUMNS[c]["width"] for c in cols]
    total_w = float(sum(weights))
    table = ax.table(cellText=body, colLabels=headers_for(cols),
                     colWidths=[w / total_w for w in weights],
                     cellLoc="center", bbox=[0, 0, 1, 1])
    table.auto_set_font_size(False)
    table.set_fontsize(8 if len(cols) > 6 else 9)

    last = len(body)
    for (r, c), cellobj in table.get_celld().items():
        cellobj.set_edgecolor("#DDDDDD")
        cellobj.set_linewidth(0.6)
        if r == 0:
            cellobj.set_facecolor(EMBER_HEX)
            cellobj.set_text_props(color="white", fontweight="bold")
        elif r == last:
            cellobj.set_facecolor("#F0F0F0")
            cellobj.set_text_props(fontweight="bold")
        elif r % 2 == 0:
            cellobj.set_facecolor("#FAFAFA")
        else:
            cellobj.set_facecolor("white")
        if r > 0 and c >= 2:
            cellobj.set_text_props(ha="right")
            cellobj.PAD = 0.04
        elif r > 0 and c == 1:
            cellobj.set_text_props(ha="left")
            cellobj.PAD = 0.03

    # Footer: text with the flame mark to its right, mirroring the PDF.
    import matplotlib.image as mpimg
    fig.text(0.014, 0.055 / fig_h, brand.FOOTER_TEXT,
             fontsize=8.5, color="#666666", va="bottom", style="italic")
    try:
        logo = mpimg.imread(brand.flame_stream(), format="png")
        mark_h = 0.16 / fig_h
        mark_w = mark_h * (fig_h / fig_w) * (logo.shape[1] / logo.shape[0])
        ax_logo = fig.add_axes([0.014 + 0.145, 0.030 / fig_h, mark_w, mark_h])
        ax_logo.imshow(logo)
        ax_logo.axis("off")
    except Exception as e:                      # never fail a report over a logo
        print(f"   note: could not draw the mark on the PNG ({e})")

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor="white", bbox_inches="tight", pad_inches=0.12)
    plt.close(fig)
    return buf.getvalue()


BUILDERS = {
    "xlsx": build_xlsx, "pdf": build_pdf, "csv": build_csv,
    "docx": build_docx, "txt": build_txt, "png": build_png,
}


def parse_formats(raw: str) -> list:
    """Turn the stored comma-separated preference into a clean ordered list."""
    wanted = {f.strip().lower() for f in (raw or "").split(",") if f.strip()}
    ordered = [f for f in SUPPORTED if f in wanted]
    return ordered or ["xlsx"]          # never send an email with no report


def build_all(church_name, rows, period_label, formats, columns=None):
    """Returns [(filename, bytes), ...] for every requested format."""
    safe_name = "".join(ch if ch.isalnum() or ch in " -_" else "" for ch in church_name)
    safe_name = safe_name.strip().replace(" ", "_") or "church"
    stamp = period_label.replace(" ", "_").replace("|", "-").replace("/", "-")
    out = []
    for fmt in formats:
        builder = BUILDERS.get(fmt)
        if not builder:
            continue
        data = builder(church_name, rows, period_label, columns)
        out.append((f"{safe_name}_livestream_report_{stamp}.{EXTENSION[fmt]}", data))
    return out
